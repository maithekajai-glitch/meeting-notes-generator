from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet


def export_markdown(notes):
    return notes.encode("utf-8")


def export_docx(notes):
    document = Document()

    document.add_heading("Meeting Notes", level=1)

    for line in notes.split("\n"):
        document.add_paragraph(line)

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output


def export_pdf(notes):
    output = BytesIO()

    doc = SimpleDocTemplate(output)

    styles = getSampleStyleSheet()

    story = []

    for line in notes.split("\n"):

        story.append(Paragraph(line.replace("\n", "<br/>"), styles["BodyText"]))

    doc.build(story)

    output.seek(0)

    return output